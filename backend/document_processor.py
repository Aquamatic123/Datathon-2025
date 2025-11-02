"""
Document Processor
Extracts text from various document formats (HTML, PDF, DOCX, TXT)
"""

import io
from typing import Optional, Dict
from bs4 import BeautifulSoup
import PyPDF2
import docx
import pdfplumber


class DocumentProcessor:
    """Extract text from various document formats"""
    
    @staticmethod
    def process_document(content: bytes, content_type: str, filename: str) -> Dict[str, any]:
        """
        Main entry point for document processing
        
        Args:
            content: Raw bytes of the document
            content_type: MIME type (e.g., 'text/html', 'application/pdf')
            filename: Original filename
            
        Returns:
            Dictionary with extracted text and metadata
        """
        # Determine file type
        file_extension = filename.lower().split('.')[-1]
        
        # Route to appropriate processor
        if content_type == 'text/html' or file_extension in ['html', 'htm']:
            text = DocumentProcessor._process_html(content)
        elif content_type == 'application/pdf' or file_extension == 'pdf':
            text = DocumentProcessor._process_pdf(content)
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword'] or file_extension in ['docx', 'doc']:
            text = DocumentProcessor._process_docx(content)
        elif content_type == 'text/plain' or file_extension == 'txt':
            text = DocumentProcessor._process_text(content)
        else:
            raise ValueError(f"Unsupported file type: {content_type} / {file_extension}")
        
        return {
            'text': text,
            'original_filename': filename,
            'content_type': content_type,
            'length': len(text),
            'word_count': len(text.split())
        }
    
    @staticmethod
    def _process_html(content: bytes) -> str:
        """
        Extract text from HTML content
        Uses your custom HTML parsing logic for directives/laws
        """
        try:
            # Decode HTML content
            html_str = content.decode('utf-8')
            
            # Use lxml parser if available, otherwise html.parser
            parser = 'html.parser'
            try:
                import lxml
                parser = 'lxml'
            except ImportError:
                pass
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_str, parser)
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            # Extract text
            plain_text = soup.get_text()
            
            # Decode HTML entities (e.g., &eacute; to é)
            import html
            plain_text = html.unescape(plain_text)
            
            # Post-conversion cleaning
            # Remove excess whitespace and newlines
            import re
            plain_text = re.sub(r'(\n\s*)+\n', '\n\n', plain_text).strip()
            
            # Additional cleanup: remove multiple spaces
            plain_text = re.sub(r' +', ' ', plain_text)
            
            return plain_text
            
        except Exception as e:
            raise ValueError(f"Error processing HTML: {str(e)}")
    
    @staticmethod
    def _process_pdf(content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            # Try pdfplumber first (better for complex PDFs)
            pdf_file = io.BytesIO(content)
            text_parts = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if text_parts:
                return '\n\n'.join(text_parts)
            
            # Fallback to PyPDF2
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            if not text_parts:
                raise ValueError("No text could be extracted from PDF")
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")
    
    @staticmethod
    def _process_docx(content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")
    
    @staticmethod
    def _process_text(content: bytes) -> str:
        """Extract text from plain text content"""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any supported encoding")


# Example usage
if __name__ == "__main__":
    # Test with a sample HTML string
    sample_html = b"""
    <!DOCTYPE html>
    <html>
    <head><title>Clean Energy Act 2024</title></head>
    <body>
        <h1>Clean Energy Advancement Act 2024</h1>
        <p>Effective Date: January 15, 2024</p>
        <p>Jurisdiction: United States</p>
        <p>This act provides $50B in subsidies for electric vehicle manufacturers...</p>
    </body>
    </html>
    """
    
    processor = DocumentProcessor()
    result = processor.process_document(sample_html, 'text/html', 'test.html')
    print("Extracted text:")
    print(result['text'])
    print(f"\nWord count: {result['word_count']}")
